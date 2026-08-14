from docx import Document

for fname, label in [('/home/ubuntu/app/Semih_Kilic_CV.docx', 'CV'), ('/home/ubuntu/app/Semih_Kilic_CoverLetter.docx', 'COVER LETTER')]:
    try:
        doc = Document(fname)
        print(f'=== {label} ===')
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text)
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    print(' | '.join(cells))
        print()
    except Exception as e:
        print(f'{label} ERROR: {e}')
